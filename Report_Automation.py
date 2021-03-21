import docx
import datetime

# This is a dictionary that converts the number of a month to its name.
date_dict = {
    1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr",
    5: "May", 6: "June", 7: "July", 8: "Aug",
    9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec"
}

# Today's date
date = datetime.date.today()

# Reading in the daily report from a .txt file
with open("reference.txt", "r") as file:
    text = file.read()

# Initialize a new .doc document
document = docx.Document()
# Enter the date
document.add_paragraph(f'Daily Report - {date_dict[date.month]} {date.day}, {date.year}')
# Enter the daily report text
document.add_paragraph(text)
# Save the document with the correct name
document.save(f"{date_dict[date.month]} {date.day}, {date.year}.docx")
